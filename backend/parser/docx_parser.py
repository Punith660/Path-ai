from __future__ import annotations

import zipfile
from io import BytesIO

from docx import Document

_MAX_UNCOMPRESSED_TOTAL = 500 * 1024 * 1024  # 500 MB
_MAX_COMPRESSION_RATIO = 100


def _validate_docx_archive(file_bytes: bytes) -> None:
    """Verify DOCX ZIP structure and guard against ZIP bombs before parsing."""
    with zipfile.ZipFile(BytesIO(file_bytes)) as zf:
        names = set(zf.namelist())
        if "[Content_Types].xml" not in names or "word/document.xml" not in names:
            raise ValueError("Invalid or corrupted DOCX file.")

        total_uncompressed = 0
        for info in zf.infolist():
            # Reject if compression ratio exceeds 100:1 (bomb indicator)
            compressed = info.compress_size
            uncompressed = info.file_size
            if compressed > 0 and uncompressed / compressed > _MAX_COMPRESSION_RATIO:
                raise ValueError("Invalid or corrupted DOCX file.")
            total_uncompressed += uncompressed
            if total_uncompressed > _MAX_UNCOMPRESSED_TOTAL:
                raise ValueError("Invalid or corrupted DOCX file.")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX paragraphs and tables into one string."""
    _validate_docx_archive(file_bytes)
    try:
        document = Document(BytesIO(file_bytes))
    except Exception as exc:
        raise ValueError("Invalid or corrupted DOCX file.") from exc

    chunks: list[str] = []

    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    chunks.extend(paragraphs)

    for table in document.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                chunks.append(" | ".join(row_cells))

    return "\n\n".join(chunks)
